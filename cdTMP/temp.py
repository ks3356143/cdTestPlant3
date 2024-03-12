from ninja_extra import api_controller, route, ControllerBase
from django.db import transaction
from pathlib import Path
from django.shortcuts import get_object_or_404
# 文档处理相关库
from docxtpl import DocxTemplate
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.text.run import CT_R
from docx.oxml.shape import CT_Picture
from docx.parts.image import ImagePart
from docx.text.run import Run
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

@route.get('/smDocument', url_name='create-smDocument')
@transaction.atomic
def create_smDocument(self, id: int):
    """生成最后说明文档"""
    # 获取项目对象
    project_obj = get_object_or_404(Project, id=id)
    # 首先第二层模版所需变量
    member = project_obj.member[0] if len(project_obj.member) > 0 else project_obj.duty_person
    context = {'name': project_obj.name, 'is_JD': False, 'ident': project_obj.ident, 'sec_title': "公开",
                'duty_person': project_obj.duty_person, 'member': member}
    if project_obj.report_type == '9':
        context['is_JD'] = True
    # 提取第一轮测试中源代码 - 用户标识
    round_1 = project_obj.pField.filter(key='0').first()
    duty_so = round_1.rdField.filter(type='SO').first()
    if not duty_so:
        return ChenResponse(code=400, status=400, message="未找到第一轮测试中源代码被测件请添加")
    context['user_ident'] = duty_so.ref

    # ~~~~~~定义模版路径~~~~~~~
    ## 定义说明最后模版
    sm_template_file = Path.cwd() / 'media' / 'form_template' / 'products' / '测试说明.docx'
    # docx生成后给docxtpl文件路径
    sm_to_tpl_file = Path.cwd() / 'media' / 'temp' / '测试说明.docx'
    # 生成最后文档路径
    sm_seitai_final_file = Path.cwd() / 'media' / 'final_seitai' / '测试说明.docx'
    # output_dir根路径
    output_files_path = Path.cwd() / 'media' / 'output_dir'
    # 获取可能使用的被复制的docx文件
    dg_copied_files = []
    sm_copied_files = []
    for file in output_files_path.iterdir():
        if file.is_file():
            if file.suffix == '.docx':
                dg_copied_files.append(file)
        elif file.is_dir():
            if file.stem == 'sm':
                for f in file.iterdir():
                    if f.suffix == '.docx':
                        sm_copied_files.append(f)
    # 打开文档找到所有的“域”
    doc = Document(sm_template_file.as_posix())
    body = doc.element.body
    sdt_element_list = body.xpath('./w:sdt')
    # 找到sdt域的名称 -> 为了对应output_dir文件 / 储存所有output_dir图片
    area_name_list = []
    image_part_list = []
    # 遍历所有控件 -> 放入area_name_list【这里准备提取公共代码】
    for sdt_ele in sdt_element_list:
        for elem in sdt_ele.iterchildren():
            # 获取“域”的名称
            if elem.tag.endswith('sdtPr'):
                for el in elem.getchildren():
                    if el.tag.endswith('alias'):
                        if len(el.attrib.values()) > 0:
                            area_name = el.attrib.values()[0]
                            area_name_list.append(area_name)
            # 开始替换里面的“域”
            if elem.tag.endswith('sdtContent'):
                elem.clear()
                if len(area_name_list) > 0:
                    area_pop_name = area_name_list.pop(0)
                    # 取到“域名称”，这里先去找media/output_dir/sm下文件，然后找media/output下文件
                    copied_file_path = ""
                    for file in sm_copied_files:
                        if file.stem == area_pop_name:
                            copied_file_path = file
                    # 这里判断是否copied_file_path没取到文件，然后遍历output_dir下文件
                    if not copied_file_path:
                        for file in dg_copied_files:
                            if file.stem == area_pop_name:
                                copied_file_path = file
                    # 找到所需的文件，将其数据复制到对应area_name的“域”
                    if copied_file_path:
                        doc_copied = Document(copied_file_path)
                        copied_element_list = []
                        element_list = doc_copied.element.body.inner_content_elements
                        for elet in element_list:
                            if isinstance(elet, CT_P):
                                copied_element_list.append(Paragraph(elet, doc_copied))
                            if isinstance(elet, CT_Tbl):
                                copied_element_list.append(Table(elet, doc_copied))
                        for para_copied in copied_element_list:
                            elem.append(para_copied._element)

                        # 下面代码就是将图片全部提取到image_part_list，以便后续插入
                        doc_copied = Document(copied_file_path)  # 需要重新获取否则namespace错误
                        copied_body = doc_copied.element.body
                        img_node_list = copied_body.xpath('.//pic:pic')
                        if not img_node_list:
                            pass
                        else:
                            for img_node in img_node_list:
                                img: CT_Picture = img_node
                                # 根据节点找到图片的关联id
                                embed = img.xpath('.//a:blip/@r:embed')[0]
                                # 这里得到ImagePart -> 马上要给新文档添加
                                related_part: ImagePart = doc_copied.part.related_parts[embed]
                                # doc_copied.part.related_parts是一个字典
                                image_part_list.append(related_part)

    # 新文档添加sdt_element_list的图片
    graph_node_list = body.xpath('.//pic:pic')
    img_count = 0
    if len(graph_node_list) == len(image_part_list):
        for graph_node in graph_node_list:
            image_run_node = getParentRunNode(graph_node)
            image_run_node.clear()
            copied_bytes_io = BytesIO(image_part_list[img_count].image.blob)
            r_element = Run(image_run_node, doc)
            inline_shape = r_element.add_picture(copied_bytes_io)
            # 设置图片位置尺寸
            source_width = inline_shape.width
            inline_shape.width = Cm(12)
            inline_shape.height = int(inline_shape.height * (inline_shape.width / source_width))
            # 设置图片所在段落居中对齐
            r_element.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            img_count += 1
    else:
        return ChenResponse(code=400, status=400, message='注意模版里面有自定义图片，请删除后重试!!!')

    try:
        doc.save(sm_to_tpl_file)
    except PermissionError as e:
        return ChenResponse(status=400, code=400, message="注意你打开了生成的文档，请关闭后再试，{0}".format(e))

    doc = DocxTemplate(sm_to_tpl_file)
    start = time.time()
    doc.render(context)  # 耗时最长，TODO:异步任务处理？或前端等待？
    end = time.time()
    print('渲染耗时：', end - start)
    try:
        doc.save(sm_seitai_final_file)
        return ChenResponse(status=200, code=200, message="最终大纲生成成功！")
    except PermissionError as e:
        return ChenResponse(status=400, code=400, message="模版文件已打开，请关闭后再试，{0}".format(e))
